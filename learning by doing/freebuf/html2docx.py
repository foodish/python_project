from docx import Document
from bs4 import BeautifulSoup
from urllib.request import Request, urlopen, urlretrieve

url = "http://freebuf.com/news/94263.html"

document = Document()
data = urlopen(url)
soup = BeautifulSoup(data, 'lxml')
article = soup.find('div', class_='articlecontent')

# content = soup.find('div', id_='contenttxt')
# title = article.find(class_='title').find('h2').get_text()
title = article.find('h2').get_text()
content = article.find_all('p')

print(title)
# print(content)
for i in content:
	if i.img:
		pic_name = ''
		print(i.img.attrs['data-original'])
		if 'gif' in i.img.attrs['data-original']:
			pic_name = 'temp.gif'
		elif 'png' in i.img.attrs['data-original']:
			pic_name = 'temp.png'
		elif 'jpg' in i.img.attrs['data-original']:
			pic_name = 'temp.jpg'
		else:
			pic_name = 'temp.jpeg'
		urlretrieve(i.img.attrs['data-original'], filename=pic_name)
		document.add_picture(pic_name)
	if i.string:
		print(i.string.encode('gbk', 'ignore'))
		document.add_paragraph(i.string)

document.save(title + '.docx')
print('success to create a document')
